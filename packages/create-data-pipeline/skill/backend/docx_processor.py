"""
Word DOCX 文件处理器 - 提取 Word 文档中的文本和表格
"""

import logging
from pathlib import Path
from typing import List, Dict, Any
from .file_processor import ProcessedDocument

logger = logging.getLogger(__name__)


class DocxProcessor:
    """Word DOCX 文档处理器"""

    def __init__(self):
        self.extract_images = False

    def process(self, file_path: str) -> ProcessedDocument:
        """处理 Word DOCX 文件"""
        try:
            from docx import Document
        except ImportError:
            return ProcessedDocument(
                content="",
                errors=["缺少依赖: python-docx。请运行: pip install python-docx"]
            )

        path = Path(file_path)
        errors = []
        content_parts = []
        tables = []
        images = []

        try:
            doc = Document(str(path))

            # 提取元数据
            metadata = {
                "source": str(path),
                "type": "docx",
                "author": doc.core_properties.author or "",
                "title": doc.core_properties.title or "",
                "subject": doc.core_properties.subject or "",
                "created": str(doc.core_properties.created) if doc.core_properties.created else "",
            }

            # 提取段落文本
            for para in doc.paragraphs:
                if para.text.strip():
                    content_parts.append(para.text)

            # 提取表格
            try:
                tables.extend(self._extract_tables(doc))
            except Exception as e:
                errors.append(f"表格提取失败: {str(e)}")

            # 提取图片
            if self.extract_images:
                try:
                    images.extend(self._extract_images(doc, path))
                except Exception as e:
                    errors.append(f"图片提取失败: {str(e)}")

            content = "\\n\\n".join(content_parts)

            return ProcessedDocument(
                content=content,
                metadata=metadata,
                tables=tables,
                images=images,
                errors=errors
            )

        except Exception as e:
            logger.error(f"DOCX 处理失败: {e}")
            return ProcessedDocument(
                content="",
                errors=[f"DOCX 处理失败: {str(e)}"]
            )

    def _extract_tables(self, doc) -> List[str]:
        """从 Word 文档中提取表格"""
        tables = []

        for table_num, table in enumerate(doc.tables):
            try:
                markdown_table = self._table_to_markdown(table)
                tables.append(f"\\n### 表格 {table_num + 1}\\n{markdown_table}")
            except Exception as e:
                logger.warning(f"表格 {table_num + 1} 提取失败: {e}")

        return tables

    def _extract_images(self, doc, docx_path: Path) -> List[Dict[str, str]]:
        """从 Word 文档中提取图片信息"""
        images = []

        try:
            for rel in doc.part.rels.values():
                if rel.target_ref.endswith((".png", ".jpg", ".jpeg", ".gif")):
                    try:
                        images.append({
                            "name": Path(rel.target_ref).name,
                            "type": rel.target_ref.split(".")[-1],
                        })
                    except Exception as e:
                        logger.warning(f"图片信息获取失败: {e}")
        except Exception as e:
            logger.warning(f"图片列表获取失败: {e}")

        return images

    def _table_to_markdown(self, table) -> str:
        """将 Word 表格转换为 Markdown 格式"""
        if not table.rows:
            return ""

        markdown_lines = []

        for row in table.rows:
            row_data = []
            for cell in row.cells:
                text = cell.text.strip()
                row_data.append(text)

            markdown_lines.append("| " + " | ".join(row_data) + " |")

        if markdown_lines:
            # 添加表头分隔线
            header_cols = len(table.rows[0].cells) if table.rows else 0
            markdown_lines.insert(1, "|" + "|".join(["---"] * header_cols) + "|")

        return "\\n".join(markdown_lines)
